#!/usr/bin/env python3
from __future__ import annotations

import argparse
import subprocess
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


JOURNAL_AUTHOR_LINE = "PARITOSH DWIVEDI and ANINDITA KUNDU"
JOURNAL_AFFILIATION = (
    "Vellore Institute of Technology, Vellore, India "
    "(e-mail: paritosh.dwivedi2024@vitstudent.ac.in; anindita.kundu@vit.ac.in)"
)
JOURNAL_CORRESPONDING = (
    "Corresponding author: Anindita Kundu "
    "(e-mail: anindita.kundu@vit.ac.in)."
)

PATENT_TITLE = (
    "System and method for GPU-accelerated multilevel Monte Carlo based "
    "network congestion propagation analysis with uncertainty quantification"
)
PATENT_FIELD = (
    "Computer networks, stochastic modeling, GPU computing, uncertainty "
    "quantification, performance engineering, and simulation-driven decision "
    "support for congestion analysis."
)
PATENT_PRIOR_ART = [
    [
        "1",
        "Multilevel Monte Carlo Path Simulation",
        "Michael B. Giles",
        "Hierarchical Monte Carlo estimator with level-wise variance "
        "reduction for stochastic differential equations",
        "Established the telescoping-sum MLMC estimator and strong "
        "theoretical complexity guarantees; however, it does not target "
        "network congestion models and does not exploit GPU hardware or "
        "graph-coupled dynamics.",
    ],
    [
        "2",
        "GPU Accelerated Monte Carlo Simulation of the 2D and 3D Ising Model",
        "Tobias Preis, Peter Virnau, Wolfgang Paul, Johannes J. Schneider",
        "Massively parallel GPU Monte Carlo kernels for large simulation batches",
        "Demonstrates the practical runtime advantage of GPU-based sampling; "
        "however, it remains a single-level Monte Carlo approach and is not "
        "designed for uncertainty-aware network performance estimation.",
    ],
    [
        "3",
        "Stochastic Networks",
        "Frank Kelly and Elena Yudovina",
        "Analytical stochastic network and queueing theory",
        "Provides rigorous foundations for stochastic network modeling and "
        "congestion behavior; however, the treatment is primarily analytical "
        "and does not provide a scalable GPU-accelerated uncertainty "
        "quantification workflow for large network topologies.",
    ],
    [
        "4",
        "Monte Carlo Methods in Financial Engineering",
        "Paul Glasserman",
        "Classical single-level Monte Carlo for stochastic systems",
        "Offers a standard Monte Carlo baseline and statistical methodology; "
        "however, the computational cost remains high for tight accuracy "
        "requirements and does not incorporate coupled graph diffusion or "
        "GPU-MLMC execution.",
    ],
]
PATENT_SUMMARY = [
    "Modern computer networks operate under stochastic traffic arrivals, "
    "service variability, topology heterogeneity, and burst-driven "
    "congestion, so deterministic or steady-state models miss the "
    "uncertainty bands required for practical provisioning decisions.",
    "Classical Monte Carlo can estimate the relevant distributions but "
    "becomes computationally expensive for tight error targets, while "
    "analytical queueing results are difficult to extend to graph-coupled, "
    "data-driven network scenarios.",
    "Existing GPU Monte Carlo literature primarily accelerates single-level "
    "sampling, and existing MLMC literature is typically developed for "
    "finance or generic SDEs rather than congestion propagation on network graphs.",
    "The novelty of the disclosed invention is the integration of a coupled "
    "congestion propagation stochastic differential equation with a "
    "GPU-accelerated multilevel Monte Carlo estimator, enabling correlated "
    "uncertainty quantification across nodes while reducing computational "
    "work relative to single-level Monte Carlo.",
    "The disclosed workflow further introduces adaptive sample allocation "
    "and a crossover-aware evaluation strategy that identifies when the "
    "multilevel method becomes more efficient than a same-hardware GPU "
    "Monte Carlo baseline.",
]
PATENT_OBJECTIVES = [
    "To provide a scalable computational framework for estimating network "
    "congestion metrics together with confidence intervals and tail-risk "
    "measures under stochastic uncertainty.",
    "To reduce the computational cost of simulation-based network "
    "uncertainty quantification by combining multilevel Monte Carlo variance "
    "reduction with GPU parallel execution.",
    "To model congestion propagation across adjacent network nodes using a "
    "coupled stochastic formulation rather than treating all queues independently.",
    "To support practical validation on synthetic and real-world Internet "
    "topology data so that the method can be used for congestion analysis, "
    "provisioning studies, and performance planning.",
]
PATENT_WORKING_PRINCIPLE = (
    "The invention models queue or congestion state evolution using "
    "stochastic differential equations. A hierarchy of discretization levels "
    "is generated, where coarse and fine paths share coupled random "
    "increments. Instead of estimating the quantity of interest only on the "
    "finest grid, the method computes a telescoping multilevel estimator "
    "composed of one coarse estimate plus correction terms across levels. "
    "Variance decays across levels, so more samples are allocated to cheaper "
    "coarse levels and fewer to expensive fine levels. The complete "
    "simulation workload is executed in parallel on the GPU, where each "
    "thread or tensor batch evaluates many stochastic paths simultaneously. "
    "For coupled propagation, the node state update includes degree-normalized "
    "influence from neighboring nodes through the adjacency matrix, which "
    "allows the method to quantify correlated congestion spread and "
    "uncertainty over the network."
)
PATENT_DESCRIPTION = [
    "The disclosed invention comprises a software-implemented system and "
    "method for uncertainty-aware network congestion analysis. A network "
    "topology is first constructed from either synthetic graph generators or "
    "empirical datasets such as CAIDA. Each node is assigned stochastic "
    "traffic and service parameters, and the congestion state is represented "
    "either by a per-node queue SDE or by a coupled congestion propagation "
    "SDE of the form:",
    "dC_i(t) = (sum_j alpha_ij C_j(t) - beta_i C_i(t)) dt + sigma_i dW_i(t)",
    "where the coefficients encode graph connectivity, damping, and noise intensity.",
    "For numerical simulation, Euler-Maruyama discretization is used at "
    "multiple time resolutions. The multilevel estimator computes "
    "expectations via level differences P_l - P_(l-1) with shared randomness "
    "so that the variance of each correction term remains small. An adaptive "
    "allocation rule determines how many paths to simulate at each level "
    "according to the measured variance-to-cost ratio. This produces lower "
    "overall cost than single-level Monte Carlo for the same target RMS error.",
    "The implementation is mapped to GPU hardware using parallel path "
    "simulation. Fine and coarse paths are executed in batches, reductions "
    "are performed on-device, and matrix-vector operations required by the "
    "coupled propagation model are evaluated directly on the GPU. The "
    "resulting output includes mean congestion, confidence intervals, and "
    "tail statistics such as high-percentile delay estimates. Experimental "
    "evidence in the associated project shows substantial runtime reduction "
    "and work reduction relative to a same-hardware GPU single-level baseline.",
    "Optional drawing/photograph placeholder: A future revision may insert "
    "a block diagram showing the workflow 'Topology/Data Input -> Coupled "
    "SDE Model -> MLMC Level Allocation -> GPU Parallel Execution -> "
    "Uncertainty Metrics'.",
]
PATENT_VALIDATION = (
    "The project validates the invention through convergence studies, GPU "
    "speed benchmarking, uncertainty quantification experiments, and "
    "real-topology evaluation. The multilevel method achieves the expected "
    "O(epsilon^-2) complexity trend and empirically reduces computational "
    "work by up to 257.72x relative to a same-GPU single-level Monte Carlo "
    "baseline at tight accuracy targets. Runtime speedup reaches up to "
    "12.91x under identical hardware and discretization settings. The "
    "implementation also demonstrates high path throughput on GPU hardware "
    "and produces accurate confidence intervals and tail-delay estimates on "
    "CAIDA AS-rel2 topology subgraphs, with reported tail-risk error below "
    "0.2% versus a tight reference Monte Carlo run. These results support "
    "both technical feasibility and lab-scale validation of the disclosed method."
)
PATENT_PROTECTION = [
    "The end-to-end method of combining multilevel Monte Carlo with "
    "GPU-parallel simulation for network congestion uncertainty quantification.",
    "The use of coupled graph-based congestion propagation SDE dynamics "
    "within a multilevel Monte Carlo estimator to obtain correlated "
    "uncertainty estimates across neighboring nodes.",
    "The adaptive level-wise sample allocation and execution strategy that "
    "reduces runtime and computational work relative to single-level GPU "
    "Monte Carlo for a target estimation accuracy.",
    "The software workflow that integrates topology ingestion, stochastic "
    "state evolution, GPU execution, and uncertainty/tail-risk reporting "
    "into a unified congestion analysis pipeline.",
]
TRL_MEANINGS = [
    "Basic principles observed",
    "Technology concept formulated",
    "Experimental proof of concept",
    "Technology validated in a lab",
    "Technology validated in relevant environment",
    "Technology demonstrated in relevant environment",
    "System prototype demonstration in operational environment",
    "System complete and qualified",
    "Actual system proven in operational environment",
]


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, *, bold: bool = False, size: int = 10,
                  align: WD_ALIGN_PARAGRAPH = WD_ALIGN_PARAGRAPH.LEFT) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = align
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "Times New Roman"
    for extra in paragraph.runs[1:]:
        extra.font.size = Pt(size)


def set_table_borders(table) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "8")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "000000")


def set_table_cell_margins(table, *, top: int = 70, start: int = 70,
                           bottom: int = 70, end: int = 70) -> None:
    tbl_pr = table._tbl.tblPr
    cell_mar = tbl_pr.first_child_found_in("w:tblCellMar")
    if cell_mar is None:
        cell_mar = OxmlElement("w:tblCellMar")
        tbl_pr.append(cell_mar)
    for name, value in (
        ("top", top),
        ("start", start),
        ("bottom", bottom),
        ("end", end),
    ):
        element = cell_mar.find(qn(f"w:{name}"))
        if element is None:
            element = OxmlElement(f"w:{name}")
            cell_mar.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")


def apply_column_widths(table, widths_cm: list[float]) -> None:
    for row in table.rows:
        for idx, width in enumerate(widths_cm):
            row.cells[idx].width = Cm(width)


def add_header_block(document: Document) -> None:
    para = document.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = para.add_run("© VIT IPR&TTCELL")
    run.font.size = Pt(8)
    run.font.name = "Times New Roman"

    table = document.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table)
    set_table_cell_margins(table)
    apply_column_widths(table, [3.3, 6.5, 4.8])

    row = table.rows[0]
    set_cell_text(row.cells[0], "VIT", bold=True, size=11, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(
        row.cells[1],
        "Invention Disclosure Format (IDF)-B",
        bold=True,
        size=11,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    row.cells[2].text = ""
    lines = [
        ("Document No.: ", "02-IPR-R003"),
        ("Issue No/Date: ", "2/01.02.2024"),
        ("Amd. No/Date: ", "0/00.00.0000"),
    ]
    for label, value in lines:
        para = row.cells[2].add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r1 = para.add_run(label)
        r1.bold = True
        r1.font.size = Pt(10)
        r1.font.name = "Times New Roman"
        r2 = para.add_run(value)
        r2.font.size = Pt(10)
        r2.font.name = "Times New Roman"
    row.cells[2].paragraphs[0].text = ""
    for cell in row.cells:
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    document.add_paragraph()


def add_labelled_section(document: Document, heading: str, body: str) -> None:
    para = document.add_paragraph()
    run = para.add_run(heading)
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(11)
    body_run = para.add_run(body)
    body_run.font.name = "Times New Roman"
    body_run.font.size = Pt(11)


def add_bullets(document: Document, heading: str, items: list[str]) -> None:
    heading_para = document.add_paragraph()
    run = heading_para.add_run(heading)
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(11)
    for item in items:
        para = document.add_paragraph(style="List Number")
        para.paragraph_format.left_indent = Cm(0.4)
        run = para.add_run(item)
        run.font.name = "Times New Roman"
        run.font.size = Pt(11)


def add_patent_prior_art(document: Document) -> None:
    heading = document.add_paragraph()
    run = heading.add_run(
        "3. Prior Patents and Publications from literature "
        "(provide a table summarizing the prior art)"
    )
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(11)

    table = document.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table)
    set_table_cell_margins(table, top=50, start=55, bottom=50, end=55)
    apply_column_widths(table, [0.8, 3.0, 2.7, 3.0, 3.1])

    headers = [
        "Sl No",
        "Title",
        "Inventors / Authors",
        "Technology used",
        "Merits and Challenges",
    ]
    for cell, header in zip(table.rows[0].cells, headers):
        set_cell_text(cell, header, bold=True, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_shading(cell, "D9E2F3")

    for row_values in PATENT_PRIOR_ART:
        row = table.add_row()
        for idx, value in enumerate(row_values):
            align = WD_ALIGN_PARAGRAPH.CENTER if idx == 0 else WD_ALIGN_PARAGRAPH.LEFT
            set_cell_text(row.cells[idx], value, size=8, align=align)


def add_trl_table(document: Document) -> None:
    heading = document.add_paragraph()
    run = heading.add_run(
        "10. What is the technology readiness level of your invention? "
        "(Tick the appropriate TRL)"
    )
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(11)

    table = document.add_table(rows=4, cols=10)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table)
    set_table_cell_margins(table, top=40, start=45, bottom=40, end=45)
    apply_column_widths(table, [0.9] + [1.65] * 9)

    top = table.rows[0].cells
    top[0].text = ""
    set_cell_text(top[1].merge(top[3]), "Research", bold=True, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(top[4].merge(top[6]), "Development", bold=True, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(top[7].merge(top[9]), "Deployment", bold=True, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)

    second = table.rows[1].cells
    set_cell_text(second[0], "", size=8, align=WD_ALIGN_PARAGRAPH.CENTER)
    for idx in range(1, 10):
        set_cell_text(second[idx], f"TRL {idx}", bold=True, size=8, align=WD_ALIGN_PARAGRAPH.CENTER)

    third = table.rows[2].cells
    set_cell_text(third[0], "Selection", bold=True, size=8, align=WD_ALIGN_PARAGRAPH.CENTER)
    for idx in range(1, 10):
        mark = "✓" if idx == 4 else ""
        set_cell_text(third[idx], mark, bold=True, size=11, align=WD_ALIGN_PARAGRAPH.CENTER)

    fourth = table.rows[3].cells
    set_cell_text(fourth[0], "Meaning", bold=True, size=8, align=WD_ALIGN_PARAGRAPH.CENTER)
    for idx, meaning in enumerate(TRL_MEANINGS, start=1):
        set_cell_text(fourth[idx], meaning, size=7, align=WD_ALIGN_PARAGRAPH.LEFT)


def set_document_defaults(document: Document, *, margin_cm: float = 2.3) -> None:
    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(11)

    for section in document.sections:
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(margin_cm)
        section.bottom_margin = Cm(margin_cm)
        section.left_margin = Cm(margin_cm)
        section.right_margin = Cm(margin_cm)


def get_pdf_page_size_cm(pdf_path: Path) -> tuple[float, float]:
    info = subprocess.run(
        ["pdfinfo", str(pdf_path)],
        check=True,
        capture_output=True,
        text=True,
    )
    for line in info.stdout.splitlines():
        if line.startswith("Page size:"):
            parts = line.split()
            width_pts = float(parts[2])
            height_pts = float(parts[4])
            return width_pts * 2.54 / 72.0, height_pts * 2.54 / 72.0
    raise RuntimeError(f"Could not determine page size for {pdf_path}")


def generate_patent_docx(output_path: Path) -> None:
    document = Document()
    set_document_defaults(document, margin_cm=1.25)
    add_header_block(document)
    add_labelled_section(document, "1. Title of the invention: ", PATENT_TITLE)
    add_labelled_section(document, "2. Field /Area of invention: ", PATENT_FIELD)
    add_patent_prior_art(document)
    add_bullets(document, "4. Summary and background of the invention (Address the gap / Novelty)", PATENT_SUMMARY)
    add_bullets(document, "5. Objective(s) of Invention", PATENT_OBJECTIVES)
    add_labelled_section(document, "6. Working principle of the invent (in brief): ", PATENT_WORKING_PRINCIPLE)

    heading = document.add_paragraph()
    run = heading.add_run("7. Description of the invention in detail (Include drawing and or photograph as needed)")
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(11)
    for idx, paragraph_text in enumerate(PATENT_DESCRIPTION):
        para = document.add_paragraph()
        para.paragraph_format.space_after = Pt(6)
        if idx == 1:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run(paragraph_text)
            run.italic = True
        else:
            run = para.add_run(paragraph_text)
        run.font.name = "Times New Roman"
        run.font.size = Pt(11)

    add_labelled_section(document, "8. Experimental validation results: ", PATENT_VALIDATION)
    add_bullets(document, "9. What aspect(s) of the invention need(s) protection?", PATENT_PROTECTION)

    document.add_page_break()
    add_header_block(document)
    add_trl_table(document)
    end_para = document.add_paragraph()
    end_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = end_para.add_run("----------------------END OF THE DOCUMENT-----------------------------")
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(10)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)


def insert_paragraph_after(paragraph, text: str, *, style: str | None = None):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = paragraph._parent.add_paragraph()
    new_para._p.getparent().remove(new_para._p)
    new_para._p = new_p
    if style:
        new_para.style = style
    run = new_para.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(10.5)
    return new_para


def remove_paragraph(paragraph) -> None:
    element = paragraph._element
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)
    paragraph._p = paragraph._element = None


def postprocess_journal_docx(docx_path: Path) -> None:
    document = Document(docx_path)
    set_document_defaults(document, margin_cm=2.2)
    author_paragraph = None
    for paragraph in document.paragraphs:
        if JOURNAL_AUTHOR_LINE in paragraph.text:
            author_paragraph = paragraph
            break
    if author_paragraph is None:
        document.save(docx_path)
        return

    affiliation_para = insert_paragraph_after(author_paragraph, JOURNAL_AFFILIATION)
    affiliation_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    corresponding_para = insert_paragraph_after(affiliation_para, JOURNAL_CORRESPONDING)
    corresponding_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    first_bibliography = None
    for paragraph in list(document.paragraphs):
        stripped = paragraph.text.strip()
        if stripped.startswith("Dwivedi : GPU-Accelerated Multilevel Monte Carlo"):
            remove_paragraph(paragraph)
            continue
        if paragraph.style.name == "Bibliography" and first_bibliography is None:
            first_bibliography = paragraph
        if stripped == (
            "Monte Carlo methods, GPU computing, network modeling, uncertainty "
            "quantification, stochastic differential equations, queueing networks"
        ):
            paragraph.text = ""
            run = paragraph.add_run(f"Keywords: {stripped}")
            run.font.name = "Times New Roman"
            run.font.size = Pt(10.5)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

    if first_bibliography is not None:
        references_heading = OxmlElement("w:p")
        first_bibliography._p.addprevious(references_heading)
        heading_para = document.add_paragraph()
        heading_para._p.getparent().remove(heading_para._p)
        heading_para._p = references_heading
        heading_para.style = "Heading1"
        heading_run = heading_para.add_run("References")
        heading_run.font.name = "Times New Roman"
        heading_run.font.size = Pt(12)
        heading_run.bold = True

    for paragraph in document.paragraphs:
        for run in paragraph.runs:
            if run.font.name is None:
                run.font.name = "Times New Roman"
            if run.font.size is None and paragraph.style.name != "Title":
                run.font.size = Pt(10.5)

    document.save(docx_path)


def generate_journal_editable_docx(project_root: Path, output_path: Path) -> None:
    tex_path = project_root / "new" / "journal" / "ieee_access_submission" / "access_project.tex"
    bib_path = project_root / "new" / "journal" / "ieee_access_submission" / "gpuAcc.bib"
    resource_dir = project_root / "new" / "journal" / "ieee_access_submission"
    output_path.parent.mkdir(parents=True, exist_ok=True)

    subprocess.run(
        [
            "pandoc",
            str(tex_path),
            f"--resource-path={resource_dir}",
            f"--bibliography={bib_path}",
            "--citeproc",
            "-o",
            str(output_path),
        ],
        check=True,
        cwd=project_root,
    )
    postprocess_journal_docx(output_path)


def generate_journal_visual_docx(project_root: Path, pdf_path: Path, output_path: Path) -> None:
    temp_dir = project_root / "tmp" / "docx_exact_images"
    temp_dir.mkdir(parents=True, exist_ok=True)
    prefix = temp_dir / "access_page"
    subprocess.run(
        [
            "pdftoppm",
            "-jpeg",
            "-jpegopt",
            "quality=92,progressive=y,optimize=y",
            "-r",
            "170",
            str(pdf_path),
            str(prefix),
        ],
        check=True,
        cwd=project_root,
    )

    image_paths = sorted(temp_dir.glob("access_page-*.jpg"))
    if not image_paths:
        raise RuntimeError(f"No rendered images produced from {pdf_path}")

    document = Document()
    page_width_cm, page_height_cm = get_pdf_page_size_cm(pdf_path)
    for section in document.sections:
        section.page_width = Cm(page_width_cm)
        section.page_height = Cm(page_height_cm)
        section.top_margin = Cm(0)
        section.bottom_margin = Cm(0)
        section.left_margin = Cm(0)
        section.right_margin = Cm(0)
    normal = document.styles["Normal"]
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.line_spacing = 1.0

    section = document.sections[0]
    for image_path in image_paths:
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run()
        run.add_picture(str(image_path), width=section.page_width)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)


def main() -> None:
    parser = argparse.ArgumentParser(description="Generate editable DOCX deliverables.")
    parser.add_argument(
        "--project-root",
        type=Path,
        default=Path.cwd(),
        help="Project root containing the new/patent and new/journal outputs.",
    )
    args = parser.parse_args()
    project_root = args.project_root.resolve()

    patent_docx = project_root / "new" / "patent" / "format_b_filled.docx"
    journal_visual_docx = project_root / "new" / "journal" / "ieee_access_submission" / "access_project.docx"
    journal_editable_docx = (
        project_root / "new" / "journal" / "ieee_access_submission" / "access_project_editable.docx"
    )
    journal_pdf = project_root / "new" / "journal" / "ieee_access_submission" / "access_project.pdf"

    generate_patent_docx(patent_docx)
    generate_journal_editable_docx(project_root, journal_editable_docx)
    generate_journal_visual_docx(project_root, journal_pdf, journal_visual_docx)


if __name__ == "__main__":
    main()
