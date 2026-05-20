#!/usr/bin/env python3
from __future__ import annotations

import shutil
import subprocess
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt


PAPER_FILES = [
    "access_project.pdf",
    "access_project.tex",
    "gpuAcc.bib",
    "ieeeaccess.cls",
    "IEEEtran.cls",
    "bullet.png",
    "logo.png",
    "notaglinelogo.png",
]

CODEBASE_FILES = [
    "README.md",
    "requirements.txt",
    "setup.py",
]

CODEBASE_DIRS = [
    "src",
    "datasets",
    "experiments",
    "examples",
]

RESULT_FILE_MAP = [
    ("paper/figures/cost_ratio_vs_epsilon_a100.png", "results/figures/cost_ratio_vs_epsilon_a100.png"),
    ("paper/figures/coupled_propagation_correlated_ci.png", "results/figures/coupled_propagation_correlated_ci.png"),
    ("paper/figures/empirical_slope_bars_a100.png", "results/figures/empirical_slope_bars_a100.png"),
    ("paper/figures/loglog_cost_vs_epsilon_a100.png", "results/figures/loglog_cost_vs_epsilon_a100.png"),
    ("paper/figures/mlmc_level_allocation_a100.png", "results/figures/mlmc_level_allocation_a100.png"),
    ("paper/figures/mlmc_variance_decay_a100.png", "results/figures/mlmc_variance_decay_a100.png"),
    ("paper/figures/runtime_scaling_tightest_epsilon_a100.png", "results/figures/runtime_scaling_tightest_epsilon_a100.png"),
    ("paper/figures/runtime_vs_epsilon_a100.png", "results/figures/runtime_vs_epsilon_a100.png"),
    ("paper/figures/scaling_n_vs_runtime.png", "results/figures/scaling_n_vs_runtime.png"),
    ("paper/figures/tail_delay_validation_a100.png", "results/figures/tail_delay_validation_a100.png"),
    ("results/results/benchmark_env.json", "results/benchmark_env.json"),
    ("results/scaling_results.json", "results/scaling_results.json"),
    ("results/exp5_coupled_propagation_results.json", "results/exp5_coupled_propagation_results.json"),
    ("results/results/tables/exp1_mlmc_convergence_results.json", "results/tables/exp1_mlmc_convergence_results.json"),
    ("results/results/tables/exp2_gpu_speedup_results.json", "results/tables/exp2_gpu_speedup_results.json"),
    ("results/results/tables/exp2_network_size_scaling.csv", "results/tables/exp2_network_size_scaling.csv"),
    ("results/results/tables/exp2_sample_size_scaling.csv", "results/tables/exp2_sample_size_scaling.csv"),
    ("results/results/tables/exp2b_accuracy_cost.csv", "results/tables/exp2b_accuracy_cost.csv"),
    ("results/results/tables/exp3_uncertainty_quantification_results.json", "results/tables/exp3_uncertainty_quantification_results.json"),
    ("results/results/tables/exp4_realworld_validation_results.json", "results/tables/exp4_realworld_validation_results.json"),
]


def reset_dir(path: Path) -> None:
    if path.exists():
        shutil.rmtree(path)
    path.mkdir(parents=True, exist_ok=True)


def copy_file(src: Path, dst: Path) -> None:
    dst.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(src, dst)


def copy_tree(src: Path, dst: Path) -> None:
    shutil.copytree(
        src,
        dst,
        dirs_exist_ok=True,
        ignore=shutil.ignore_patterns(
            "__pycache__",
            ".DS_Store",
            ".gitkeep",
            ".pytest_cache",
            "*.pyc",
        ),
    )


def add_code_block(document: Document, block: str) -> None:
    paragraph = document.add_paragraph(style="No Spacing")
    paragraph.paragraph_format.left_indent = Cm(0.45)
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(8)
    for line in block.strip("\n").splitlines():
        run = paragraph.add_run(line)
        run.font.name = "Courier New"
        run.font.size = Pt(8.5)
        paragraph.add_run("\n")


def build_run_guide(codebase_dir: Path) -> None:
    docx_path = codebase_dir / "run_instructions.docx"
    pdf_path = codebase_dir / "run_instructions.pdf"

    document = Document()
    normal = document.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(11)

    for section in document.sections:
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(1.8)
        section.bottom_margin = Cm(1.8)
        section.left_margin = Cm(1.8)
        section.right_margin = Cm(1.8)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("GPU-Acc-Net-Prop-Congestion-Multi-Monte-Carlo Run Guide")
    run.bold = True
    run.font.size = Pt(18)

    intro = document.add_paragraph()
    intro.add_run(
        "This guide accompanies the curated codebase copy inside submission/codebase/. "
        "It lists the minimum steps to run the project, a verified smoke-test path, "
        "and the GPU-specific requirements for the heavier experiments."
    )

    for heading, paragraph_text in [
        (
            "What Is Included",
            "The submission codebase contains src/, datasets/, experiments/, examples/, "
            "README.md, requirements.txt, and setup.py. It excludes notebooks, tests, "
            "scratch files, and unrelated packaging artifacts.",
        ),
        (
            "Verified CPU Smoke Test",
            "The command python3 examples/basic_simulation.py was run successfully on "
            "March 29, 2026 in this workspace. That path does not require CUDA and is "
            "the recommended first check after copying the submission folder to another machine.",
        ),
    ]:
        document.add_heading(heading, level=2)
        document.add_paragraph(paragraph_text)

    add_code_block(
        document,
        """
cd codebase
python3 -m venv venv
source venv/bin/activate
pip install --upgrade pip
pip install numpy scipy networkx matplotlib pandas
python examples/basic_simulation.py
        """,
    )

    document.add_heading("Full Project Install", level=2)
    document.add_paragraph(
        "Use the full dependency install only on a machine where the required "
        "scientific and CUDA dependencies are available. This is the path to use "
        "when you want the paper experiments and GPU benchmarks."
    )
    add_code_block(
        document,
        """
cd codebase
python3 -m venv venv
source venv/bin/activate
pip install --upgrade pip
pip install -r requirements.txt
pip install -e .
        """,
    )

    document.add_heading("GPU Notes", level=2)
    document.add_paragraph(
        "GPU-specific runs require an NVIDIA GPU, a working CUDA toolchain, and Python "
        "packages such as pycuda and numba. The current machine did not have those GPU "
        "Python packages installed, so only the CPU smoke test was validated end-to-end here."
    )

    document.add_heading("Suggested Commands", level=2)
    add_code_block(
        document,
        """
# Basic example (verified)
python examples/basic_simulation.py

# MLMC convergence study
python experiments/exp1_mlmc_convergence.py --output-dir results/exp1

# GPU speedup benchmark (CUDA machine)
pip install joblib
python experiments/exp2_gpu_speedup.py --output-dir results/exp2

# Real-world validation
python experiments/exp4_realworld_validation.py --output-dir results/exp4
        """,
    )

    document.add_heading("Dataset Notes", level=2)
    document.add_paragraph(
        "The copied datasets/ folder includes the dataset loader code and synthetic "
        "generators, but not large external CAIDA or MAWI downloads. Several scripts "
        "fall back to synthetic data if those external datasets are absent, so the code "
        "remains runnable for demonstration and smoke tests."
    )

    document.add_heading("Outputs", level=2)
    document.add_paragraph(
        "By default, experiment outputs are written under a results/ folder in the "
        "working directory. The paper figures already used for the IEEE Access manuscript "
        "are stored separately in the paper package and do not need to be regenerated "
        "unless you are rerunning experiments."
    )

    document.save(docx_path)
    subprocess.run(
        [
            "/opt/homebrew/bin/soffice",
            "-env:UserInstallation=file:///tmp/lo_codex_submission_guide",
            "--headless",
            "--convert-to",
            "pdf",
            "--outdir",
            str(codebase_dir),
            str(docx_path),
        ],
        check=True,
    )


def main() -> None:
    project_root = Path(__file__).resolve().parent.parent
    submission_root = project_root / "submission"
    paper_src = project_root / "new" / "journal" / "ieee_access_submission"
    patent_src = project_root / "new" / "patent" / "format_b_filled.docx"

    paper_dst = submission_root / "ieee_access_paper"
    patent_dst = submission_root / "patent"
    codebase_dst = submission_root / "codebase"

    reset_dir(paper_dst)
    reset_dir(patent_dst)
    reset_dir(codebase_dst)

    for filename in PAPER_FILES:
        copy_file(paper_src / filename, paper_dst / filename)
    copy_tree(paper_src / "figures", paper_dst / "figures")

    copy_file(patent_src, patent_dst / patent_src.name)

    for filename in CODEBASE_FILES:
        copy_file(project_root / filename, codebase_dst / filename)
    for dirname in CODEBASE_DIRS:
        copy_tree(project_root / dirname, codebase_dst / dirname)
    for src_rel, dst_rel in RESULT_FILE_MAP:
        copy_file(project_root / src_rel, codebase_dst / dst_rel)

    (codebase_dst / "results" / "README.txt").write_text(
        "This results folder contains the paper-relevant generated outputs copied from the\n"
        "main project workspace. It includes the figures used in the IEEE Access paper and\n"
        "the main JSON/CSV result tables that back the reported experiments.\n",
        encoding="utf-8",
    )

    build_run_guide(codebase_dst)


if __name__ == "__main__":
    main()
